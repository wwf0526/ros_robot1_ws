from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "ROS2连续体机器人框架_GPT项目上下文.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.style = doc.styles["No Spacing"]
    r = p.add_run(text)
    r.font.name = "DejaVu Sans Mono"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    r.font.size = Pt(8.5)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    footer.add_run(" 页")


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)
    add_page_number(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    r = title.add_run("ROS 2 连续体机器人控制与 MuJoCo 仿真框架")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("面向 GPT 深度交流与协同开发的项目上下文文档")
    sr.font.size = Pt(15)
    sr.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()
    add_table(doc, ["项目项", "内容"], [
        ["项目目录", "/home/wangwenfeng/ros_robot1_ws"],
        ["GitHub", "https://github.com/wwf0526/ros_robot1_ws.git"],
        ["ROS 2 版本", "Jazzy"],
        ["主要语言", "Python 3 + ROS 2 接口定义 + MuJoCo MJCF XML"],
        ["机器人类型", "两段、六绳驱动、双 IMU 连续体机器人"],
        ["文档日期", "2026-08-14"],
        ["文档目的", "将本文件与源码一起提供给 GPT，使其准确理解架构、接口、现状和开发约束"],
    ], [4, 12])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("重要：本文描述的是当前源码状态，明确区分已实现能力、实验性能力与待完成内容。").bold = True
    doc.add_page_break()

    add_heading(doc, "1. 如何把本文件交给 GPT 使用", 1)
    doc.add_paragraph("建议把本 Word 文档与需要修改的源码文件一起上传。若只能上传一个文件，可先上传本文件，让 GPT 建立整体认识，再按任务补充对应包的源代码。")
    add_numbered(doc, [
        "先要求 GPT 阅读全文并复述其对硬件、数据流、接口和当前缺陷的理解，不要立即改代码。",
        "告诉 GPT 本次任务只涉及哪个包、是否允许改变 ROS 接口、是否连接真实电机。",
        "要求 GPT 在修改前列出影响范围，修改后给出构建、仿真和实机分级验证方法。",
        "任何可能驱动真实电机的测试，都必须先使用仿真、mock CAN 或禁用控制输出验证。",
    ])
    add_heading(doc, "1.1 可直接复制给 GPT 的总提示词", 2)
    add_code(doc, """你现在作为本项目的 ROS 2 连续体机器人高级开发助手。请先完整阅读本项目上下文，不要凭包名猜测实现状态。

项目是一套 ROS 2 Jazzy 两段六绳连续体机器人框架，包含双 CAN 电机驱动、双 IMU、PCC 状态估计、MPC、消息服务和 MuJoCo 运动学可视化。

协作规则：
1. 修改前说明数据流、受影响的包、话题、消息和参数。
2. 保持 robot_interfaces 的兼容性；如必须修改接口，要同时更新发布者、订阅者、构建文件和文档。
3. 不得默认连接或驱动真实电机。涉及 /motor/command、/motor/command_array、回零或急停服务时，先提供无硬件验证方案。
4. 单位必须明确：电机角度 deg、速度 rad/s、绳长 mm、PCC 长度 m、曲率 1/m、姿态根据字段使用 deg 或 rad。
5. 真实串联顺序为 base -> section2 -> section1 -> tip；section1 靠近末端，section2 靠近基座。
6. 线缆映射和几何参数以 robot_calibration.yaml 为唯一事实来源，避免在多个节点重复硬编码。
7. 修改后至少检查 Python 语法、colcon build、单元测试和 ROS 接口连通性；实机测试必须单独列为人工步骤。
8. 当前 MuJoCo 节点主要通过写 qpos + mj_forward() 做运动学显示，不能误称为完整动力学闭环。

请先回答：你对系统边界、核心数据流、当前实现程度、主要风险和本次任务影响范围的理解。""")

    add_heading(doc, "2. 项目目标与系统边界", 1)
    doc.add_paragraph("本项目为深腔作业场景下的两段线驱连续体机器人提供底层控制与研究平台。目标是把真实硬件采集、电机控制、连续体状态估计、解析模型、优化控制和仿真显示连接到统一的 ROS 2 数据总线上。")
    add_bullets(doc, [
        "硬件层：两路 CAN 总线、6 台 MS42DDC 类电机、双 IMU 串口设备。",
        "接口层：统一的电机、绳长、连续体状态、目标、安全和残差消息，以及零点/回零/急停服务。",
        "估计层：电机编码器角度转换为绳长，结合双 IMU 和 PCC 模型产生机器人状态。",
        "控制层：简化线性控制器与单步 QP/OSQP 控制器输出六电机目标。",
        "仿真层：把 ContinuumState 中的 PCC 状态映射到 MuJoCo 多关节近似模型进行运动学显示。",
        "研究扩展：后续可加入系统辨识、摩擦/迟滞/重力残差、RL/PINNs 补偿和真正动力学仿真。",
    ])
    add_heading(doc, "2.1 不应误解的边界", 2)
    add_bullets(doc, [
        "当前没有证据表明所有节点已经在完整实机闭环中联合验证。",
        "当前 MuJoCo 不是电机—绳索—柔性体的完整动力学仿真，而是 PCC 状态驱动的运动学可视化。",
        "MPC 中的 B 矩阵是根据六绳几何给出的初始线性估计，尚未由实机辨识替换。",
        "安全状态主要集成在状态估计节点中，尚无独立且完整的 safety supervisor。",
    ])

    add_heading(doc, "3. 机器人结构与关键约定", 1)
    add_table(doc, ["属性", "约定"], [
        ["结构", "2 个串联 section，每段 3 根等效驱动绳，共 6 电机"],
        ["串联顺序", "base -> section2（基端）-> section1（末端）-> tip"],
        ["电机分组", "section1: motor 1,3,5；section2: motor 2,4,6"],
        ["物理绳角", "t1=300°, t2=240°, t3=180°, t4=120°, t5=60°, t6=0°"],
        ["PCC section1 槽位", "[DL1,DL2,DL3] = [tendon5,tendon3,tendon1]"],
        ["PCC section2 槽位", "[DL1,DL2,DL3] = [tendon4,tendon2,tendon6]"],
        ["IMU 映射", "section1 -> imu1；section2 -> imu2"],
        ["当前名义长度", "section1 L0=200 mm；section2 L0=140 mm（由盘厚和逐段间隙计算）"],
        ["绳半径", "当前标定 tendon_radius_mm=55"],
    ])
    add_heading(doc, "3.1 单位约定", 2)
    add_table(doc, ["量", "单位", "备注"], [
        ["电机目标/反馈位置", "degree", "MotorCommand.target_deg、MotorState.position_deg"],
        ["电机速度", "rad/s", "MotorCommand.speed_rad_s"],
        ["绳长", "mm", "TendonState 与 ContinuumState 的 tendon_length_mm"],
        ["IMU 姿态", "degree", "ContinuumState 的 imu*_roll/pitch/yaw"],
        ["PCC 弯曲角和方向角", "rad", "theta_rad、phi_rad"],
        ["PCC 曲率", "1/m", "kappa_1pm"],
        ["PCC 坐标和长度", "m", "px/py/pz、arc_length、center_length、l0"],
    ])

    add_heading(doc, "4. 总体软件架构与数据流", 1)
    add_code(doc, """真实硬件链路
  双 IMU 串口
      -> imu_driver
      -> /imu1/data_raw, /imu2/data_raw -----------+
                                                    |
  CAN0/CAN1 + 6 电机                              v
      <-> motor_can_driver                    continuum_state_estimator
      -> /motor/state -------------------------->  |  电机角度 -> 绳长
      -> /motor/emergency_stop_active ----------> |  IMU 四元数 -> 欧拉角
                                                   |  PCC 正运动学
                                                   v
                  /continuum/tendon_state
                  /continuum/state
                  /continuum/safety_state
                           |
              +------------+------------------+
              v                               v
       continuum_mpc                    continuum_mujoco_sim
       订阅目标与安全状态                订阅 /continuum/state
       -> /motor/command_array           写入 qpos + mj_forward
              |                          显示机器人几何姿态
              v
       motor_can_driver -> 真实电机

残差旁路
  /continuum/state -> model_residual_node -> /continuum/model_residual
                                              -> state_estimator 安全判断""")

    add_heading(doc, "5. ROS 2 包清单与职责", 1)
    package_rows = [
        ["robot_interfaces", "ament_cmake", "自定义 msg/srv，是所有包之间的接口契约", "可构建"],
        ["motor_can_driver", "ament_python", "双 CAN 电机协议、命令、状态、零点、回零、急停", "可构建；需硬件验证"],
        ["imu_driver", "ament_python", "双 IMU 串口读取、解析、发布和监听工具", "可构建；需设备验证"],
        ["robot_bringup", "ament_python", "硬件、标定、MPC、安全参数集合", "可构建；缺总 launch"],
        ["continuum_model", "ament_python", "基础 PCC/曲率计算", "可构建；与新 PCC 包部分重叠"],
        ["continuum_pccmodel", "ament_python", "PCC 几何、正运动学、逆运动学、变换", "可构建"],
        ["continuum_state_estimator", "ament_python", "电机+IMU+PCC 状态估计和安全状态", "可构建"],
        ["continuum_mpc", "ament_python", "线性控制器与 OSQP 单步 QP 控制器", "可构建；模型待辨识"],
        ["continuum_mujoco_sim", "ament_python", "MJCF 生成、PCC 状态可视化", "当前安装阶段失败"],
    ]
    add_table(doc, ["包", "构建类型", "职责", "当前状态"], package_rows)

    add_heading(doc, "6. 节点、话题与服务", 1)
    node_rows = [
        ["motor_node", "motor_can_driver", "/motor/command, /motor/command_array", "/motor/state, /motor/emergency_stop_active", "50 Hz 状态发布"],
        ["dual_imu_serial_node", "imu_driver", "串口字节流", "/imu1/data_raw, /imu2/data_raw", "默认 200 Hz"],
        ["continuum_state_estimator", "continuum_state_estimator", "/motor/state, 两路 IMU, 急停, model_residual", "tendon/state/safety", "50 Hz"],
        ["imu_zero_check_node", "continuum_state_estimator", "两路 IMU", "/continuum/zero_check", "默认 10 Hz"],
        ["model_residual_node", "continuum_state_estimator", "/continuum/state", "/continuum/model_residual", "模型与 IMU 姿态残差"],
        ["linear_mpc", "continuum_mpc", "state, target, safety", "/motor/command_array", "10 Hz 简化控制"],
        ["qp_mpc_node", "continuum_mpc", "state, target, safety", "/motor/command_array", "默认 5 Hz OSQP"],
        ["mujoco_pcc_viewer", "continuum_mujoco_sim", "/continuum/state", "MuJoCo GUI", "运动学显示"],
        ["generate_mujoco_model", "continuum_mujoco_sim", "robot_calibration.yaml", "生成 MJCF XML", "配置驱动模型生成"],
        ["pcc_section_demo", "continuum_mujoco_sim", "无", "/pcc/section_state", "演示数据源"],
    ]
    add_table(doc, ["节点", "包", "主要输入", "主要输出", "说明"], node_rows)
    add_heading(doc, "6.1 电机服务", 2)
    add_table(doc, ["服务", "类型", "作用"], [
        ["/motor/set_zero", "SetZero", "记录指定或全部电机当前位置为零点"],
        ["/motor/home_motors", "HomeMotors", "指定或全部电机按给定速度回零"],
        ["/motor/emergency_stop", "EmergencyStop", "触发软件急停"],
        ["/motor/clear_estop", "ClearEmergencyStop", "解除软件急停"],
    ])

    add_heading(doc, "7. 自定义消息与服务接口", 1)
    interface_rows = [
        ["MotorCommand", "motor_id:uint8; mode:uint8; target_deg:float32; speed_rad_s:float32"],
        ["MotorCommandArray", "header; commands:MotorCommand[]"],
        ["MotorState", "header; motor_id; raw_deg; position_deg; speed_rad_s; reached; channel"],
        ["TendonState", "header; tendon_length_mm[]; motor_position_deg[]; tendon_slack[]"],
        ["ContinuumTarget", "两段 roll/pitch 目标(deg); enable; mode"],
        ["SafetyState", "safe_to_control; slack; IMU timeout; motor limit; residual; estop; status"],
        ["ZeroCheckState", "两路 IMU RPY(deg)、容差、各 IMU/整机调平标志、status"],
        ["ModelResidual", "两段模型/IMU roll,pitch；四项残差；valid；status"],
        ["ContinuumState", "两段曲率、双 IMU、绳长、松绳标志、两段 PCC 输入/参数/位姿、整体末端位姿"],
    ]
    add_table(doc, ["接口", "核心字段"], interface_rows)
    doc.add_paragraph("ContinuumState 是当前系统最重要的状态接口。修改它会影响接口生成、状态估计、MPC、残差节点和 MuJoCo viewer，应视为跨包 API 变更。")
    service_rows = [
        ["SetZero", "motor_ids:int32[]", "success, message"],
        ["HomeMotors", "motor_ids:int32[]; speed_rad_s", "success, message"],
        ["EmergencyStop", "空", "success, message"],
        ["ClearEmergencyStop", "空", "success, message"],
    ]
    add_table(doc, ["服务类型", "请求", "响应"], service_rows)

    add_heading(doc, "8. 配置文件及事实来源", 1)
    add_table(doc, ["文件", "用途", "关键内容"], [
        ["hardware_params.yaml", "真实硬件", "串口、波特率、CAN 接口、两路通道、电机 ID、话题、速度、零点文件"],
        ["robot_calibration.yaml", "机器人几何与标定", "电机符号、卷线轮半径、绳角、段长度、IMU 映射、PCC 几何与链顺序"],
        ["mpc_paras.yaml", "控制器参数", "频率、增量限制、位置限制、姿态增益、电机分段"],
        ["safety_params.yaml", "安全阈值", "IMU/电机超时、限位余量、松绳、残差、急停开关"],
        ["zero_offsets.json", "零点持久化", "电机 1..6 的零点偏移"],
    ])
    add_bullets(doc, [
        "理想设计中 robot_calibration.yaml 应成为几何和映射的唯一事实来源。",
        "当前部分参数仍硬编码在节点中，且 state_estimator 与 motor_node 使用绝对路径，降低了可移植性。",
        "mpc_paras.yaml 与实际节点参数名并非完全一致，需要统一。",
        "safety_params.yaml 尚未被完整、统一地加载，部分阈值仍直接写在 Python 中。",
    ])

    add_heading(doc, "9. 核心算法说明", 1)
    add_heading(doc, "9.1 电机角度到绳长", 2)
    doc.add_paragraph("状态估计器读取 MotorState.position_deg，结合电机方向 sign、卷线轮半径 spool_radius_mm 和零点，计算每根绳的长度变化。必须统一正负号约定：当前配置注释定义 position_deg 增大为收绳时 sign=1。")
    add_code(doc, "概念公式：delta_length_mm = sign * spool_radius_mm * radians(position_deg - zero_deg)\n实际实现与消息语义必须共同验证正负方向。")
    add_heading(doc, "9.2 PCC 模型", 2)
    add_bullets(doc, [
        "continuum_pccmodel 根据每段三根绳长变化计算弯曲角 theta、弯曲方向 phi、曲率 kappa 和中心线长度。",
        "正运动学先计算 section2，再串联 section1：T_end = T_section2 * T_section1。",
        "逆运动学使用 scipy.optimize.minimize，根据目标位姿求绳长变化。",
        "equal_dl_tol_m 用于判断三根绳近似等长、当前段不弯曲的退化情况。",
    ])
    add_heading(doc, "9.3 状态估计", 2)
    add_bullets(doc, [
        "IMU 四元数转换为 roll/pitch/yaw，并减去安装偏置。",
        "电机角度转换成六根绳长，检测松绳状态。",
        "基础模型产生 section1/section2 的 kx、ky。",
        "新 PCC 模型填充每段和整体末端位姿字段。",
        "根据 IMU 超时、松绳、残差、限位和急停生成 SafetyState。",
    ])
    add_heading(doc, "9.4 控制器", 2)
    add_bullets(doc, [
        "linear_mpc：使用六绳角度构造 2×6 几何矩阵，当前更接近简化增量控制，而非完整时域 MPC。",
        "qp_mpc_node：状态 x=[imu1_roll, imu1_pitch, imu2_roll, imu2_pitch]，控制 u 为 6 个电机角度增量。",
        "QP 目标同时最小化下一步姿态误差和电机动作，约束单步增量及绝对电机角度。",
        "当前 A=I，B 为基于 60° 绳布局的初始估计并乘 0.05，后续应由辨识模型替换。",
        "只有收到 state、target、safety 且 safe_to_control=true 时才允许输出命令。",
    ])
    add_heading(doc, "9.5 MuJoCo", 2)
    add_bullets(doc, [
        "generate_mujoco_model.py 从 robot_calibration.yaml 的盘厚和逐段间隙生成 MJCF，使仿真几何与 PCC 长度一致。",
        "mujoco_pcc_viewer.py 订阅 ContinuumState，把每段弯曲均分到若干 bend_x/bend_y/slide_z 关节。",
        "节点直接写 qpos 并调用 mujoco.mj_forward()，因此没有电机力矩、绳张力、接触、柔性材料或闭环动力学响应。",
    ])

    add_heading(doc, "10. 当前启动与构建方式", 1)
    add_code(doc, """cd /home/wangwenfeng/ros_robot1_ws
source /opt/ros/jazzy/setup.bash
colcon build
source install/setup.bash

# 已有单独启动入口
ros2 launch imu_driver imu.launch.py
ros2 launch motor_can_driver motor.launch.py

# 状态估计与控制节点
ros2 run continuum_state_estimator state_estimator_node
ros2 run continuum_state_estimator imu_zero_check_node
ros2 run continuum_state_estimator model_residual_node
ros2 run continuum_mpc qp_mpc_node

# MuJoCo
ros2 run continuum_mujoco_sim generate_mujoco_model
ros2 run continuum_mujoco_sim mujoco_pcc_viewer""")
    p = doc.add_paragraph()
    p.add_run("警告：").bold = True
    p.add_run("当前完整 colcon build 会在 continuum_mujoco_sim 安装资源时失败；以上命令是目标用法，不代表当前全部可直接运行。启动 motor_node 前必须确认 CAN 配置和机械安全。")

    add_heading(doc, "11. 已验证状态", 1)
    add_table(doc, ["检查项", "结果", "说明"], [
        ["Python 语法编译", "通过", "src 下 Python 文件 compileall 无语法错误"],
        ["ROS 2 Jazzy 完整构建", "失败", "continuum_mujoco_sim 的 setup.py 把 assets/__pycache__ 目录当文件复制"],
        ["跳过 MuJoCo 构建", "通过", "其余 8 个包全部构建完成"],
        ["真实 CAN/电机", "未在本次审计验证", "不能仅凭构建成功认为硬件安全可用"],
        ["双 IMU 串口", "未在本次审计验证", "需要真实设备和数据帧测试"],
        ["闭环 MPC", "未完整验证", "需要仿真、辨识和低速实机分阶段验证"],
        ["MuJoCo GUI", "存在实现", "当前属于运动学可视化，且包安装有缺陷"],
    ])

    add_heading(doc, "12. 已知缺陷与遗漏（按优先级）", 1)
    add_heading(doc, "P0：阻碍可靠构建或安全运行", 2)
    add_bullets(doc, [
        "修复 continuum_mujoco_sim/setup.py 的资源 glob，只安装普通文件或明确列出扩展名。",
        "增加 robot_bringup 总 launch，支持 hardware/simulation、安全开关和控制器选择。",
        "禁止默认驱动真实电机；建议增加 dry_run/mock_can/use_sim 参数。",
        "统一加载 safety_params.yaml，避免安全阈值在代码和配置中分裂。",
        "消除绝对路径，改用 ament_index_python 定位 share/config。",
    ])
    add_heading(doc, "P1：影响复现和持续开发", 2)
    add_bullets(doc, [
        "在 package.xml/依赖清单中补齐 python-can、pyserial、PyYAML、NumPy、SciPy、OSQP、MuJoCo。",
        "增加 requirements.txt 或 rosdep 可解析的依赖说明和一键环境安装文档。",
        "统一 mpc_paras.yaml 与节点 declare_parameter 的参数名和层级。",
        "增加协议、IMU 解析、PCC、MPC 约束、MuJoCo XML 和 ROS 图集成测试。",
        "README 更新为当前 9 包架构，补充安装、构建、启动、硬件连接和故障排查。",
        "明确 continuum_model 与 continuum_pccmodel 的职责，避免重复模型逐渐漂移。",
    ])
    add_heading(doc, "P2：工程质量与扩展", 2)
    add_bullets(doc, [
        "增加 CI：colcon build、colcon test、flake8、接口兼容检查。",
        "补充 TF/坐标系发布、URDF/robot_state_publisher 或明确完全采用 MJCF 的理由。",
        "引入时间戳同步、QoS 配置、传感器失联恢复、CAN 重连和诊断消息。",
        "加入 rosbag 数据记录、实验元数据、标定版本和可重复数据集流程。",
        "将 B 矩阵替换为实验辨识模型，逐步扩展到多步 MPC、迟滞/摩擦残差和约束软化。",
        "构建真实动力学 MuJoCo 模型：actuator、tendon、sensor、接触和闭环仿真接口。",
    ])

    add_heading(doc, "13. Git 与 GitHub 当前状态", 1)
    add_table(doc, ["项目", "状态"], [
        ["远端 main", "1fae0f1 (v2.0)"],
        ["本地分支", "pcc_model_v1 @ a9bd71a"],
        ["本地领先远端", "4 个提交：PCC FK/IK、状态估计集成、MuJoCo、PCC 长度计算"],
        ["未提交修改", "5 个已修改文件；generate_mujoco_model.py 与 generated XML 尚未跟踪"],
        ["严重仓库卫生问题", ".venv_mujoco 已被 Git 跟踪，约 3811 个文件；推送前必须清理"],
    ])
    doc.add_paragraph("因此 GitHub 当前不是本地完整框架的镜像。不要直接 push 当前分支；先将虚拟环境从 Git 索引/历史中移除、补充 .gitignore、修复构建并整理提交。")

    add_heading(doc, "14. 建议的目标目录结构", 1)
    add_code(doc, """ros_robot1_ws/
├── README.md
├── requirements.txt                 # 或明确的 rosdep/venv 安装方案
├── .github/workflows/ci.yml
├── docs/
│   ├── architecture.md
│   ├── hardware.md
│   └── safety.md
├── src/
│   ├── robot_interfaces/
│   ├── motor_can_driver/
│   ├── imu_driver/
│   ├── robot_bringup/
│   │   ├── launch/
│   │   │   ├── hardware.launch.py
│   │   │   ├── simulation.launch.py
│   │   │   └── full_system.launch.py
│   │   └── config/
│   ├── continuum_model/             # 或与 pccmodel 合并
│   ├── continuum_pccmodel/
│   ├── continuum_state_estimator/
│   ├── continuum_mpc/
│   └── continuum_mujoco_sim/
├── test/
└── tools/

# 不应进入 Git
build/ install/ log/ .venv*/ __pycache__/ *.pyc""")

    add_heading(doc, "15. 推荐开发路线", 1)
    add_numbered(doc, [
        "仓库清理：移除被跟踪的 .venv_mujoco、备份 msg 文件和生成缓存；完善 .gitignore。",
        "恢复可构建：修复 MuJoCo setup.py；补齐依赖；保证 9 包全量 colcon build。",
        "统一配置：去掉绝对路径，统一 MPC 和 safety 参数，定义唯一几何事实来源。",
        "建立总 launch：硬件模式、仿真模式、只感知模式、只显示模式分开。",
        "增加测试：先纯函数，再节点接口，再 mock CAN/录制 IMU 数据集成测试。",
        "仿真闭环：MuJoCo 发布虚拟电机/IMU 状态，MPC 输出进入仿真执行器。",
        "系统辨识：用 rosbag 采集电机—绳长—IMU—末端位姿数据，辨识 B/残差模型。",
        "低速实机闭环：限制角度、速度、增量，配置急停和人工监护，逐步扩大工作区。",
    ])

    add_heading(doc, "16. GPT 修改代码时的检查清单", 1)
    add_bullets(doc, [
        "是否理解 section1/section2 的物理位置与计算顺序？",
        "是否保留所有消息字段的单位和语义？",
        "是否同时更新 package.xml、setup.py、launch 和 config？",
        "是否避免硬编码用户目录和设备名称？",
        "是否提供不连接电机的测试方法？",
        "是否避免让两个控制节点同时向 /motor/command_array 发布？",
        "是否在 SafetyState 不允许控制时停止输出？",
        "是否考虑 IMU/CAN 超时、消息未到达、空数组和异常数值？",
        "是否运行完整 colcon build/test，并区分构建成功与实机验证？",
        "是否说明对 GitHub 分支和未提交修改的影响？",
    ])

    add_heading(doc, "17. 面向后续对话的任务模板", 1)
    add_code(doc, """本次任务目标：<填写具体目标>
允许修改的包：<包名>
禁止修改的接口：<msg/srv/topic/API>
运行模式：<纯算法 / MuJoCo / mock CAN / 真实硬件>
硬件状态：<未连接 / 已连接但禁用 / 允许低速测试>
验收标准：<构建、测试、话题输出、控制性能>

请按以下顺序回答并执行：
1. 复述你对任务和现有数据流的理解；
2. 指出风险、单位和接口影响；
3. 给出最小修改方案；
4. 实现代码与配置；
5. 给出无硬件验证；
6. 单独列出需要人工执行的实机步骤。""")

    add_heading(doc, "18. 总结", 1)
    doc.add_paragraph("该工作空间已经形成了连续体机器人研究框架的主要骨架：底层电机和 IMU、ROS 接口、PCC 模型、状态估计、MPC 与 MuJoCo 均有代码实现。当前最大问题不是缺少模块名称，而是模块之间尚未达到可复现、可一键启动、可测试和可安全闭环的工程完成度。后续开发应优先修复构建与仓库问题，再统一配置和启动，随后建立仿真闭环与自动测试，最后进入受控的实机验证。")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
